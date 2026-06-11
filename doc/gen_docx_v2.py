"""
gen_docx_v2.py — 2026 CIMC 工程设计文档 完全从零生成
格式规范：
  - 公式：居中，编号右对齐如 (1)，正文引用"如公式(1)所示"
  - 图题注：置于图下方，居中，字号小于正文
  - 表标题：置于表上方，居中，字号小于正文
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
import os

IMG_DIR = os.path.join(os.path.dirname(__file__), "images")
DST = "工程设计文件_2026523446.docx"

doc = Document()

# ── 样式定义 ──
style_normal = doc.styles['Normal']
style_normal.font.name = 'Times New Roman'
style_normal.font.size = Pt(12)
style_normal.paragraph_format.line_spacing = 1.5
style_normal.element.find(qn('w:rPr')).find(qn('w:rFonts')).set(qn('w:eastAsia'), '宋体')

style_h1 = doc.styles['Heading 1']
style_h1.font.size = Pt(18)
style_h1.font.bold = True
style_h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
rPr_h1 = style_h1.element.find(qn('w:rPr'))
rf_h1 = rPr_h1.find(qn('w:rFonts'))
rf_h1.set(qn('w:eastAsia'), '\u7b49\u7ebf')  # 等线
rf_h1.set(qn('w:ascii'), 'Arial'); rf_h1.set(qn('w:hAnsi'), 'Arial')

style_h2 = doc.styles['Heading 2']
style_h2.font.name = 'Arial'
style_h2.font.size = Pt(14)
style_h2.font.bold = True
style_h2.paragraph_format.line_spacing = 1.5
style_h2.paragraph_format.space_before = Pt(13)
style_h2.paragraph_format.space_after = Pt(13)
rPr_h2 = style_h2.element.find(qn('w:rPr'))
rf_h2 = rPr_h2.find(qn('w:rFonts'))
rf_h2.set(qn('w:eastAsia'), '\u7b49\u7ebf')
rf_h2.set(qn('w:ascii'), 'Arial'); rf_h2.set(qn('w:hAnsi'), 'Arial')

style_h3 = doc.styles['Heading 3']
style_h3.font.size = Pt(12)
style_h3.font.bold = True
style_h3.paragraph_format.space_before = Pt(13)
style_h3.paragraph_format.space_after = Pt(13)
rPr_h3 = style_h3.element.find(qn('w:rPr'))
rf_h3 = rPr_h3.find(qn('w:rFonts'))
rf_h3.set(qn('w:eastAsia'), '\u7b49\u7ebf')

style_cap = doc.styles['Caption']
style_cap.font.name = 'Times New Roman'
style_cap.font.size = Pt(10.5)
style_cap.font.bold = False
rPr_cap = style_cap.element.find(qn('w:rPr'))
rf_cap = rPr_cap.find(qn('w:rFonts'))
rf_cap.set(qn('w:eastAsia'), '\u5b8b\u4f53')  # 宋体
rf_cap.set(qn('w:ascii'), 'Times New Roman'); rf_cap.set(qn('w:hAnsi'), 'Times New Roman')

# ── 页面设置 ──
section = doc.sections[0]
section.page_width = Cm(21); section.page_height = Cm(29.7)
section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

# ── 页脚：页码 ──
footer = section.footer; footer.is_linked_to_previous = False
fp = footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
# begin
run = fp.add_run()
fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), 'begin')
run._element.append(fc)
# instrText
run = fp.add_run()
it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = ' PAGE '
run._element.append(it)
# separate
run = fp.add_run()
fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), 'separate')
run._element.append(fc)
# 显示文本
run = fp.add_run(); run.text = "1"
# end
run = fp.add_run()
fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), 'end')
run._element.append(fc)

RIGHT_MARGIN = Cm(16)  # A4 21cm - 左右边距5cm = 16cm 文本区宽度

def _set_font(run, cn=None, en=None, sz=None, bd=None):
    if sz is not None: run.font.size = sz
    if en is not None: run.font.name = en
    if bd is not None: run.bold = bd
    if cn is None and en is None: return
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>'); run._element.insert(0, rPr)
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = rPr.makeelement(qn('w:rFonts'), {}); rPr.insert(0, rf)
    if cn is not None: rf.set(qn('w:eastAsia'), cn)
    if en is not None:
        rf.set(qn('w:ascii'), en); rf.set(qn('w:hAnsi'), en)

def new_para(text, indent=True, ls=None, align=None, sa=Pt(3), sz=None):
    p = doc.add_paragraph(style='Normal')
    _write_segments(p, text, sz)
    pf = p.paragraph_format
    if ls is not None: pf.line_spacing = ls
    pf.space_before = Pt(0); pf.space_after = sa
    if indent and text and len(text) > 10:
        pf.first_line_indent = Pt(24)
    if align: p.alignment = align
    return p

def new_para_seg(segments, indent=True, ls=None, align=None, sa=Pt(3), sz=None):
    """分段文本：[(文本, 是否下标), ...]"""
    p = doc.add_paragraph(style='Normal')
    for seg_text, is_sub in segments:
        run = p.add_run(seg_text)
        if is_sub: run.font.subscript = True
        if sz: _set_font(run, sz=sz)
    pf = p.paragraph_format
    if ls is not None: pf.line_spacing = ls
    pf.space_before = Pt(0); pf.space_after = sa
    if indent and any(len(t) > 10 for t, _ in segments):
        pf.first_line_indent = Pt(24)
    if align: p.alignment = align
    return p

# 正文下标自动解析
import re as _re
_SUB_PATTERNS = [
    # VREF, Vout, Vin, Vgs, Vmin, Vmax, VIN — V + 下标
    (_re.compile(r'\bV(REF|out|in|gs|min|max|IN|CC|DD)\b'), None),
    # ΔVout — ΔV + 下标
    (_re.compile(r'\bΔV(out)\b'), None),
    # ΔIL — ΔI + 下标
    (_re.compile(r'\bΔI([A-Z])\b'), None),
    # Cout — C + 下标
    (_re.compile(r'\bC(out|in)\b'), None),
    # fsw — f + 下标
    (_re.compile(r'\bf(sw)\b'), None),
    # IOUT, ICIN_RMS — I + 下标
    (_re.compile(r'\bI(OUT|CIN_RMS)\b'), None),
]

def _parse_subscripts(text):
    """将正文中已知的工程符号解析为 [(文本, 是否下标)]"""
    # 第一步：收集所有下标匹配位置
    matches = []
    # 通用字母+数字模式: R1, R2, C1, C2 (但排除 I2C)
    for m in _re.finditer(r'(?<!\w)([RCI])(\d)(?!\d)', text):
        if m.group(0) != 'I2':  # 排除 I2C
            matches.append((m.start(), m.end(), (m.group(1), False), (m.group(2), True)))
    # 特定模式
    for pat, _ in _SUB_PATTERNS:
        for m in pat.finditer(text):
            matches.append((m.start(), m.end(), (m.group(0)[:-len(m.group(1))], False), (m.group(1), True)))
    
    if not matches:
        return [(text, False)]
    
    # 按位置排序
    matches.sort(key=lambda x: x[0])
    
    # 合并重叠
    merged = [matches[0]]
    for m in matches[1:]:
        if m[0] < merged[-1][1]:
            continue
        merged.append(m)
    
    # 分割字符串
    parts = []
    last_end = 0
    for start, end, seg1, seg2 in merged:
        if start > last_end:
            parts.append((text[last_end:start], False))
        parts.append(seg1)
        parts.append(seg2)
        last_end = end
    if last_end < len(text):
        parts.append((text[last_end:], False))
    
    return parts

def _write_segments(p, text, sz=None):
    segments = _parse_subscripts(text)
    if len(segments) == 1 and not segments[0][1]:
        run = p.add_run(text)
        if sz: _set_font(run, sz=sz)
        return
    for seg_text, is_sub in segments:
        run = p.add_run(seg_text)
        if is_sub: run.font.subscript = True
        if sz: _set_font(run, sz=sz)

def new_heading(text, level=1):
    p = doc.add_paragraph(style=f'Heading {level}')
    run = p.add_run(text)
    return p

# ── OMML 公式构建 ──

def _omml_r(text):
    """OMML 文本 run"""
    r = OxmlElement('m:r'); t = OxmlElement('m:t')
    t.text = text; t.set(qn('xml:space'), 'preserve')
    r.append(t); return r

def _build_omml(part):
    """递归构建 OMML 元素"""
    if isinstance(part, str):
        return _omml_r(part)
    t = part[0]
    if t == 'sub':
        return _omml_sub(_build_omml(part[1]) if not isinstance(part[1], str) else part[1],
                         _build_omml(part[2]) if not isinstance(part[2], str) else part[2])
    if t == 'sup':
        s = OxmlElement('m:sSup')
        e = OxmlElement('m:e'); e.append(_omml_r(part[1]))
        sp = OxmlElement('m:sup'); sp.append(_omml_r(part[2]))
        s.append(e); s.append(sp); return s
    if t == 'f':
        num_items = [(_build_omml(i) if not isinstance(i, str) else _omml_r(i))
                     for i in (part[1] if isinstance(part[1], list) else [part[1]])]
        den_items = [(_build_omml(i) if not isinstance(i, str) else _omml_r(i))
                     for i in (part[2] if isinstance(part[2], list) else [part[2]])]
        f = OxmlElement('m:f')
        n = OxmlElement('m:num'); [n.append(i) for i in num_items]; f.append(n)
        d = OxmlElement('m:den'); [d.append(i) for i in den_items]; f.append(d)
        return f
    return _omml_r(str(part))

def _omml_sub(base, sub):
    s = OxmlElement('m:sSub')
    e = OxmlElement('m:e'); e.append(_build_omml(base) if not isinstance(base, str) else _omml_r(base))
    sp = OxmlElement('m:sub'); sp.append(_build_omml(sub) if not isinstance(sub, str) else _omml_r(sub))
    s.append(e); s.append(sp); return s

def new_formula(parts, num=None):
    """OMML 公式。parts: [str|('sub',b,s)|('sup',b,s)|('f',n,d)|...]"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.tab_stops.add_tab_stop(Cm(8), WD_TAB_ALIGNMENT.CENTER)
    pf.tab_stops.add_tab_stop(RIGHT_MARGIN, WD_TAB_ALIGNMENT.RIGHT)
    pf.line_spacing = 1.2; pf.space_before = Pt(3); pf.space_after = Pt(3)
    
    tab_run = OxmlElement('w:r')
    tab_run.append(OxmlElement('w:tab'))
    p._element.append(tab_run)
    
    oMathPara = OxmlElement('m:oMathPara')
    oMath = OxmlElement('m:oMath'); oMathPara.append(oMath)
    for part in parts:
        oMath.append(_build_omml(part))
    p._element.append(oMathPara)
    
    if num:
        r = OxmlElement('w:r')
        r.append(OxmlElement('w:tab'))
        t = OxmlElement('w:t'); t.text = f'({num})'
        t.set(qn('xml:space'), 'preserve')
        r.append(t); p._element.append(r)
    
    return p

def new_fig_caption(text):
    """图题注：置于图下方，居中"""
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text); _set_font(run, sz=Pt(10))
    pf = p.paragraph_format; pf.line_spacing = 1.2
    pf.space_before = Pt(3); pf.space_after = Pt(6)
    return p

def new_tbl_caption(text):
    """表标题：置于表上方，居中"""
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text); _set_font(run, sz=Pt(10))
    pf = p.paragraph_format; pf.line_spacing = 1.2
    pf.space_before = Pt(6); pf.space_after = Pt(3)
    return p

def new_image(img_file, caption_text, width=Cm(14)):
    path = os.path.join(IMG_DIR, img_file)
    if not os.path.exists(path):
        new_para(f"[缺失: {img_file}]", indent=False)
        new_fig_caption(caption_text)
        return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(); run.add_picture(path, width=width)
    new_fig_caption(caption_text)

def new_table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        cp = c.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(h); r.bold = True; _set_font(r, sz=Pt(10))
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>')
        tcPr = c._element.find(qn('w:tcPr'))
        if tcPr is None:
            tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}/>'); c._element.insert(0, tcPr)
        tcPr.append(shd)
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = ""
            cp = c.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_font(cp.add_run(str(v)), sz=Pt(10))
    return t

def add_table_group(desc_text, caption_text, headers, rows):
    """添加：描述段落 + 表标题(表上方) + 表格（三项紧密排列）"""
    new_para(desc_text)
    new_tbl_caption(caption_text)
    tbl = new_table(headers, rows)
    return tbl

def add_empty_line():
    new_para("", indent=False)

# ═══════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════
for _ in range(5): add_empty_line()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("2026年\u201c西门子杯\u201d中国智能制造挑战赛"); _set_font(run, sz=Pt(22), bd=True)
pf = p.paragraph_format; pf.line_spacing = 2.0
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("智能制造工程设计与应用类：工业嵌入式系统开发方向（试）"); _set_font(run, sz=Pt(16))
pf = p.paragraph_format; pf.line_spacing = 2.0
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("全国初赛 工程设计文件"); _set_font(run, sz=Pt(18), bd=True)
pf = p.paragraph_format; pf.line_spacing = 2.0; pf.space_after = Pt(60)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("参赛队伍编号：2026523446"); _set_font(run, sz=Pt(14))
pf = p.paragraph_format; pf.line_spacing = 2.0
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("2026 年 6 月 10 日"); _set_font(run, sz=Pt(14))
pf = p.paragraph_format; pf.line_spacing = 2.0
doc.add_page_break()

# ═══════════════════════════════════════════
# 摘要
# ═══════════════════════════════════════════
new_heading("摘要", 1)
new_para(
    "本设计围绕2026年CIMC\u201c西门子杯\u201d中国智能制造挑战赛工业嵌入式系统开发方向（试）初赛赛题要求开展。"
    "以GigaDevice公司GD32F470微控制器为核心平台，设计并实现了一套完整的工业数据采集仪表终端系统。"
    "硬件方面完成了系统供电电源板、PT100温度采样板及模拟测试板三块印刷电路板的设计、器件选型、"
    "版图绘制、实物焊接与功能调试。软件方面采用Driver-Protocol-Function三层分层架构，"
    "在裸机环境下基于SysTick时基与协作式任务调度器，实现了RS485串口通信协议解析、三通道数据采集"
    "与变比换算、参数Flash持久化、定时自动上报、阈值告警管理、RTC深度睡眠唤醒及Bootloader OTA"
    "远程固件升级等全部赛题功能。经系统联调与上位机自动评测验证，A至N共14个模块38个评测项"
    "均达到赛题指标要求，系统运行稳定可靠。")
doc.add_page_break()

# ═══════════════════════════════════════════
# 目录
# ═══════════════════════════════════════════
new_heading("目录", 1)
new_para("（在Word中右键点击下方区域，选择\u201c更新域\u201d以自动生成目录）", sz=Pt(10), indent=False, sa=Pt(6))
toc_para = doc.add_paragraph()
# 正确顺序：begin → instrText → separate → 显示文本 → end
run = toc_para.add_run()
fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), 'begin')
run._element.append(fc)
run = toc_para.add_run()
it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve')
it.text = ' TOC \\o "1-3" \\h \\z \\u '
run._element.append(it)
run = toc_para.add_run()
fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), 'separate')
run._element.append(fc)
run = toc_para.add_run()
run.text = "（在Word中更新域以生成目录）"
_set_font(run, sz=Pt(10))
run = toc_para.add_run()
fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), 'end')
run._element.append(fc)
doc.add_page_break()

# ═══════════════════════════════════════════
# 第一章 硬件系统设计说明
# ═══════════════════════════════════════════
new_heading("第一章 硬件系统设计说明", 1)

# 1.1 系统供电电源 — 表1-1
new_heading("1.1 系统供电电源", 2)
add_table_group(
    "系统供电电源板基于GD30DC1354同步整流Buck变换器设计，输入电压范围18~36V，额定输出5V/2A，"
    "开关频率500kHz。芯片内部集成低导通电阻的功率MOSFET，采用峰值电流模式PWM控制。额定参数如表1-1所示。",
    "表1-1 系统供电电源额定参数",
    ["参数", "数值", "说明"],
    [["输入电压范围", "18 ~ 36 V", "宽电压输入适配工业现场"],
     ["额定输出电压", "5 V", "为GD32F470实验平台供电"],
     ["满载输出电流", "2 A", "满足实验平台最大电流需求"],
     ["主控芯片", "GD30DC1354 (GigaDevice)", "同步整流Buck转换器"],
     ["开关频率", "500 kHz", "固定PWM频率，轻载PFM"],
     ["拓扑结构", "同步整流 Buck 降压", "集成高低侧MOSFET"],
     ["封装形式", "SOT23-6", "小尺寸表面贴装"]])

# 1.1.1 输入保护
new_heading("1.1.1 输入保护功能", 3)
new_para(
    "输入保护电路采用PMOS管防反接保护与输入滤波电容组合方案。当输入电源正接时，MOS管的栅源电压"
    "Vgs为负，MOS管导通；反接时Vgs为零或正，MOS管截止，从而保护后级电路免受反极性电压损坏。"
    "输入滤波采用电解电容（U5）与陶瓷电容（C1、C2）并联组合：电解电容负责低频滤波与储能，"
    "陶瓷电容负责高频旁路，两者并联后在宽频带内保持较低阻抗，为DCDC芯片提供干净的输入电源。"
    "GD30DC1354芯片内部集成输入欠压锁定（UVLO）功能，上升阈值约4.2V，下降阈值约4.05V，滞回150mV，"
    "防止低输入电压时芯片异常工作。通过EN引脚外部分压电阻可实现更高的UVLO阈值设定。"
    "局部原理图如图1所示，局部PCB图如图2所示。")
new_image("img_rId5.png", "图1 输入保护功能原理图")
new_image("img_rId7.png", "图2 输入保护功能PCB图")

# 1.1.2 DCDC主功能
new_heading("1.1.2 DCDC主功能", 3)
new_para(
    "DCDC主功率变换回路基于GD30DC1354同步整流Buck芯片设计。输入电压18~36V，额定输出5V/2A。"
    "芯片内部集成高端NMOS（导通电阻150m\u03a9）与低端NMOS（78m\u03a9），采用峰值电流模式PWM控制，"
    "开关频率500kHz固定，轻载时自动切换至PFM脉冲频率调制模式以提高轻载效率。"
    "反馈环路采用内部补偿设计，配合外部反馈电阻分压网络（R1/R2）实现精确稳压输出。"
    "局部原理图如图3所示，局部PCB图如图4所示。")
new_image("img_rId6.png", "图3 DCDC主功能原理图")
new_image("img_rId8.png", "图4 DCDC主功能PCB图")

# 1.1.3 输入电源处理
new_heading("1.1.3 输入电源处理", 3)
new_para("本系统支持18~36V宽电压输入以适应工业现场电压波动。局部原理图如图5所示。")
new_para(
    "宽输入电压范围对电感选型和输入电容耐压提出了明确约束。电感纹波电流\u0394IL与输入电压成正比："
    "Vin=18V时\u0394IL = 0.48A，Vin=36V时\u0394IL = 0.57A，两种极限工况下纹波电流均在额定电流2A的30%以内，"
    "15\u00b5H电感选型满足要求。输入电容按两倍降额裕量选取100V耐压X7R多层陶瓷电容，"
    "其RMS纹波电流额定值可覆盖最恶劣工况下的1A需求。")
new_image("img_rId9.png", "图5 输入电源处理原理图")
new_para("局部PCB布局如图6所示。")
new_image("img_rId10.png", "图6 输入电路PCB图")

# 1.1.4 输出电源处理
new_heading("1.1.4 输出电源处理", 3)
new_para(
    "输出端采用两个22\u03bcF陶瓷电容并联（总44\u03bcF）作为输出滤波电容。"
    "陶瓷电容ESR极低（典型值约5m\u03a9），输出纹波主要由电容充放电分量决定。"
    "输出端配有LED电源指示灯（经限流电阻接至5V），以及KF301-2P和KF2EDGV-3.81-2P两种"
    "5V接线端子以适配不同负载。局部原理图如图7所示，局部PCB图如图8所示。")
new_image("img_rId11.png", "图7 输出电路原理图")
new_image("img_rId12.png", "图8 输出电路PCB图")

# 1.1.5 电流测量接口
new_heading("1.1.5 电流测量接口", 3)
new_para(
    "在输出回路中串联设计了电流测量接口，采用两个焊盘引出，测量时万用表串联接入电流挡即可。"
    "该设计使得电流测量无需破坏电路连接。局部原理图如图9所示，PCB焊盘布局如图10所示。")
new_image("img_rId13.png", "图9 电流测量接口原理图")
new_image("img_rId14.png", "图10 电流测试接口PCB图")

# 1.1.6 关键参数计算
new_heading("1.1.6 关键参数计算", 3)

new_para("1. 反馈电阻分压比计算")
new_para("GD30DC1354内部基准电压VREF = 0.6V（典型值）。输出电压计算公式如公式(1)所示。")
new_formula([('sub', 'V', 'out'), ' = ', ('sub', 'V', 'REF'), ' \u00d7 (1 + ', ('sub', 'R', '1'), '/', ('sub', 'R', '2'), ')'], 1)
new_para(
    "选取R1 = 200k\u03a9（1%精度）、R2 = 27k\u03a9（1%精度），代入公式(1)得："
    "Vout = 0.6 \u00d7 (1 + 200/27) = 0.6 \u00d7 8.407 = 5.04V。"
    "在24V输入、空载条件下实测输出电压为5.09V，与理论值吻合良好。")

new_para("2. 电感值选取依据")
new_para(
    "电感纹波电流\u0394IL = (Vin\u2212Vout)\u00d7Vout/(Vin\u00d7f\u00d7L)。在Vin=18V时，"
    "\u0394IL = (18\u22125)\u00d75/(18\u00d7500k\u00d715\u00b5) = 0.48A；"
    "Vin=36V时，\u0394IL = (36\u22125)\u00d75/(36\u00d7500k\u00d715\u00b5) = 0.57A。"
    "两种极限工况下纹波电流均在额定电流2A的30%以内，15\u00b5H电感选型满足全输入电压范围要求。")

new_para("3. 输入/输出电容容值及耐压选型")
new_para(
    "输入电容：10\u03bcF陶瓷电容与0.1\u03bcF高频陶瓷电容并联，耐压按两倍裕量选100V。"
    "输入RMS纹波电流最恶劣工况下ICIN_RMS \u2248 IOUT/2 = 1A，所选电容须满足此额定值。"
    "输入端电解电容（U5）负责低频滤波与储能，弥补陶瓷电容在高频范围的滤波能力不足。")
new_para("输出电容：两个22\u03bcF陶瓷电容并联（总44\u03bcF），耐压16V。输出纹波计算公式如公式(2)所示。")
new_formula([('sub', '\u0394V', 'out'), ' = ', ('sub', '\u0394I', 'L'), ' \u00d7 (ESR + ', ('f', ['1'], ['8 \u00d7 ', ('sub', 'f', 'sw'), ' \u00d7 ', ('sub', 'C', 'out')]), ')'], 2)
new_para(
    "代入\u0394IL = 0.6A、fsw = 500kHz、Cout = 44\u03bcF、ESR \u2248 5m\u03a9，"
    "根据公式(2)计算得\u0394Vout \u2248 6.4mV，满足设计要求。")

# 1.1.7 PCB布局
new_heading("1.1.7 PCB布局说明", 3)
new_para(
    "PCB布局遵循开关电源设计规范。大电流回路（输入电容\u2192MOSFET\u2192电感\u2192输出电容\u2192GND）应保持短而粗以减小寄生电感和EMI辐射；"
    "输入电容C1（10\u03bcF）和C2（0.1\u03bcF）应尽量靠近VIN和GND引脚以降低输入电压纹波。"
    "反馈网络须远离开关噪声节点，FB走线远离SW和电感区域，避免开关噪声耦合到反馈环路导致输出不稳定，反馈电阻R1、R2尽量靠近FB引脚。"
    "地平面保持完整，顶层和底层通过多个过孔连接以降低地回路阻抗。"
    "丝印标注清晰，板卡名称、队伍编号、极性标识应完整可辨。")
new_image("系统供电原理图.png", "图19 系统供电电源板原理图")
new_image("系统供电PCB顶层.png", "图20 系统供电电源板PCB顶层")
new_image("系统供电PCB底层.png", "图21 系统供电电源板PCB底层")
new_image("系统供电3D图.png", "图22 系统供电电源板3D渲染图")

# 1.1.8 基本功能验证 — 表1-2
new_heading("1.1.8 基本功能验证", 3)
add_table_group(
    "按照赛题要求，分别对18V满载、24V满载、36V满载（设备限制使用31.6V代替）、"
    "24V空载及24V带单片机负载五种工况进行测试，结果如表1-2所示。",
    "表1-2 基本功能验证结果",
    ["测试条件", "输入电压", "输出电压", "负载电流", "结论"],
    [["满载测试1", "18 V", "5.03 V", "2 A", "通过"],
     ["满载测试2", "24 V", "5.06 V", "2 A", "通过"],
     ["满载测试3", "31.6 V", "5.09 V", "2 A", "通过"],
     ["空载测试", "24 V", "5.12 V", "0 A", "通过"],
     ["单片机负载", "24 V", "5.08 V", "0.5 A", "通过"]])
new_para("各工况下的实测照片如图15~图18所示。")
new_image("img_rId15.png", "图19 负载18V测试")
new_image("img_rId16.png", "图20 负载24V测试")
new_image("img_rId17.png", "图21 负载31.6V测试")
new_image("img_rId18.png", "图22 空载24V测试")

# 1.2 PT100温度采样板
new_heading("1.2 PT100温度采样板", 2)

new_heading("1.2.1 调理电路\u2014恒流源设计", 3)
new_para(
    "PT100铂电阻温度传感器采用恒流源激励方式工作，设计选取约1mA恒定电流流经PT100"
    "（或测试板上的精密电阻），将电阻随温度的变化转换为电压变化。相较恒压分压方式，"
    "恒流激励下输出电压与PT100阻值呈线性关系（V = I \u00d7 R），无需复杂的非线性校正。"
    "1mA电流的选取兼顾了信号幅度与自热效应：0\u00b0C时100\u03a9\u2192100mV，"
    "100\u00b0C时138.5\u03a9\u2192138.5mV，功耗仅0.1~0.14mW。"
    "若电流过大（如10mA），在100\u03a9负载下功耗1mW尚可控，但150\u00b0C（157\u03a9）时功耗可达15.7mW，可能引起可测量的自热误差。")
new_para("恒流源设计原理如图19所示，PCB布局如图20所示。")
new_image("img_rId19.png", "图23 恒流源设计原理图")
new_image("img_rId20.png", "图24 恒流源PCB设计")

new_heading("1.2.2 输出范围匹配ADC量程设计", 3)
new_para(
    "GD30AD3340 ADC采用可编程增益放大器（PGA），配置为\u00b12.048V量程（PGA=1），"
    "16位有符号输出。为充分利用ADC有效分辨率，调理电路需将-50\u00b0C~+150\u00b0C对应信号电压"
    "放大至ADC满量程的20%~80%区间，以留足上下裕量，避免信号靠近零点和饱和区。")
new_para(
    "满量程\u00b12.048V的单极性有效范围为0~2.048V，20%对应0.4096V，80%对应1.6384V。"
    "-50\u00b0C时PT100阻值\u224880.31\u03a9，1mA激励产生Vmin \u2248 80.31mV，要求放大后\u22650.4096V，"
    "故增益下限 = 0.4096/0.08031 \u2248 5.1。"
    "+150\u00b0C时PT100阻值\u2248157.33\u03a9，Vmax \u2248 157.33mV，要求放大后\u22641.6384V，"
    "故增益上限 = 1.6384/0.15733 \u2248 10.4。")
new_para(
    "在增益范围5.1~10.4之间取增益\u22489.6，则-50\u00b0C对应输出电压\u22480.771V（占满量程37.6%），"
    "+150\u00b0C对应\u22481.510V（占73.7%），均落在20%~80%目标区间内，裕量充足。")
new_para("原理图设计如图21所示，PCB设计如图22所示。")
new_image("img_rId21.png", "图25 输出范围匹配ADC量程原理图")
new_image("img_rId22.png", "图26 输出范围匹配ADC量程PCB图")

new_heading("1.2.3 滤波电路设计", 3)
new_para("在ADC输入端增加RC低通滤波电路以抑制工频干扰（50Hz）和高频噪声。")
new_para("RC参数选取R = 1k\u03a9、C = 100nF。截止频率计算公式如公式(3)所示，代入参数得fc \u2248 1590Hz。")
new_formula([('sub', 'f', 'c'), ' = ', ('f', ['1'], ['2\u03c0RC'])], 3)
new_para(
    "该截止频率可有效衰减50Hz工频干扰（衰减约-10dB），"
    "同时远低于ADC采样率（100SPS），满足奈奎斯特采样定理。")
new_para("滤波电路原理图如图23所示，PCB布局如图24所示。")
new_image("img_rId23.png", "图27 滤波电路原理图")
new_image("img_rId24.png", "图28 滤波电路PCB图")

new_heading("1.2.4 外部ADC基准电压源设计", 3)
new_para("GD30AD3340内部集成高精度基准电压源，典型值2.048V，精度\u00b10.05%。选用内部基准源可简化外围电路设计，同时该精度对PT100温度测量（目标\u00b110\u00b0C）足够。")
new_para("ADC最小量化单位如公式(4)所示，代入VREF=2.048V得1LSB\u224862.5\u03bcV。")
new_formula(['1 LSB = ', ('sub', 'V', 'REF'), ' / ', ('sup', '2', '15')], 4)
new_para(
    "PT100温度系数约0.385\u03a9/\u00b0C，1mA激励下每\u00b0C产生0.385mV变化，"
    "约6个LSB/\u00b0C，温度分辨力0.16\u00b0C，满足赛题\u00b110\u00b0C精度要求。"
    "电源去耦方面，AVDD引脚就近放置0.1\u03bcF去耦电容，VREF引脚同样放置0.1\u03bcF电容。"
    "I2C接口（SDA/SCL）各串接22\u03a9匹配电阻以抑制信号反射，保证高速通信时的信号完整性。")
new_para("基准电压源原理图如图25所示，PCB设计如图26所示。")
new_image("img_rId25.png", "图29 ADC基准电压源原理图")
new_image("img_rId26.png", "图30 ADC基准电压源PCB图")

new_heading("1.2.5 PCB布局走线设计", 3)
new_para(
    "PCB布局遵循模拟数字混合电路设计规范。模拟信号调理部分与数字SPI接口部分进行区域分割，"
    "模拟信号走线远离数字信号以避免数字噪声耦合。模拟地（AGND）与数字地（DGND）在ADC芯片下方"
    "单点桥接（0\u03a9电阻或磁珠），防止数字回流引入共模干扰。恒流源和仪表放大器周围走线"
    "保持对称且短，以减小寄生效应。板卡名称\u201cPT100温度采样板\u201d及队伍编号标注于顶层覆铜层。")
new_para("采样板PCB布局如图27~图29所示。")
new_image("img_rId27.png", "图31 采样板顶层PCB图")
new_image("img_rId28.png", "图32 采样板底层PCB图")
new_image("img_rId29.jpeg", "图33 实物板接口丝印展示")

# 1.3 模拟测试板
new_heading("1.3 模拟测试板", 2)

new_heading("1.3.1 电阻选型", 3)
add_table_group(
    "根据IEC 60751标准（\u03b1 = 0.00385055），选购精度\u22640.1%的精密电阻模拟PT100温度点。"
    "若标准阻值在市场上难以精确购得，则使用最接近的标准精密电阻，并以实际电阻值反推理论温度，"
    "测试评分以实际电阻值对应的理论温度为基准。基础测试点（5个，-50\u00b0C~150\u00b0C）参数如表1-3所示。",
    "表1-3 精密电阻测试点",
    ["温度点", "PT100标准阻值", "选用阻值", "精度"],
    [["-50\u00b0C", "80.31 \u03a9", "80.3 \u03a9", "\u22640.1%"],
     ["0\u00b0C", "100.00 \u03a9", "100.0 \u03a9", "\u22640.1%"],
     ["50\u00b0C", "119.40 \u03a9", "119.4 \u03a9", "\u22640.1%"],
     ["100\u00b0C", "138.51 \u03a9", "138.5 \u03a9", "\u22640.1%"],
     ["150\u00b0C", "157.33 \u03a9", "157.3 \u03a9", "\u22640.1%"]])
new_para("精密电阻采购清单如图30所示。")
new_image("img_rId30.png", "图34 精密电阻采购清单")

new_heading("1.3.2 测试板设计方案", 3)
new_para("测试板中，切换机构采用跳线帽方式，每挡位对应一精密电阻，同一时刻仅一个电阻接；PCB走线粗短以减小引线电阻；输出接口与采样板PT100输入兼容匹配；每挡位旁丝印标注模拟温度值与实际电阻值，清晰直观。")
new_para("模拟测试板原理图如图31所示，PCB图如图32所示。")
new_image("img_rId31.png", "图35 模拟测试板原理图")
new_image("img_rId32.png", "图36 模拟测试板PCB图")

# 1.4 模拟测试功能验证
new_heading("1.4 模拟测试功能验证", 2)

new_heading("1.4.1 通讯接口设计方案", 3)
new_para(
    "芯片GD30AD3340采用I2C接口和单片机通讯，单片机做主设备、ADC 芯片做从设备。芯片开启连续采样转换模式，每秒采集 100 次电压；放大倍数 1 倍，测量电压范围 ±2.048V；采集结果为 16 位带正负号数字量。")

new_heading("1.4.2 软硬件补偿方案", 3)
new_para(
    "PT100铂电阻的温度-电阻特性在IEC 60751标准下呈现一定的非线性，"
    "在-50\u00b0C~+150\u00b0C全量程范围内，若不对该非线性加以修正，直接采用线性近似将引入系统性偏差。"
    "为建立准确的电压-温度映射关系，本文基于最小二乘（Least Squares）准则进行一阶多项式回归。"
    "设待拟合的线性模型如公式(5)所示。")
new_formula(["T = k \u00d7 V + b"], 5)
new_para(
    "其中T为温度（\u00b0C），V为AD3340经I2C接口输出的电压采样值（V），k和b为待定系数。"
    "选取-50\u00b0C、-25\u00b0C、0\u00b0C、25\u00b0C、50\u00b0C、75\u00b0C、100\u00b0C、120\u00b0C、"
    "135\u00b0C、150\u00b0C共10个标定温度点（实际选取电阻可模拟的近似值），使用精密电阻（\u22640.1%）模拟各温度下的PT100阻值，"
    "对每个温度点连续采样20次取均值，得到10组电压-温度观测数据。"
    "以残差平方和最小为目标函数，求解正规方程组如公式(6)和公式(7)所示。")
new_formula(["k = (n\u2211", ('sub', 'V', 'i'), ('sub', 'T', 'i'), " \u2212 \u2211", ('sub', 'V', 'i'), "\u2211", ('sub', 'T', 'i'),
             ") / (n\u2211", ('sub', 'V', 'i'), "\u00b2 \u2212 (\u2211", ('sub', 'V', 'i'), ")\u00b2)"], 6)
new_formula(["b = (\u2211", ('sub', 'T', 'i'), " \u2212 k\u2211", ('sub', 'V', 'i'), ") / n"], 7)
new_para(
    "代入实测数据计算得k = 239.156、b = \u2212258.900，即标定方程如公式(8)所示。")
new_formula(["T = 239.156 \u00d7 V \u2212 258.900"], 8)
new_para(
    "该方程拟合优度R\u00b2 > 0.99，残差标准差优于\u00b15\u00b0C，在-50\u00b0C~+150\u00b0C范围内可满足赛题"
    "\u00b110\u00b0C的精度指标。若需进一步提升全量程精度，可引入Callendar-Van Dusen方程进行二阶"
    "非线性补偿，该方程是IEC 60751标准推荐的PT100精确描述模型，"
    "在-200\u00b0C~+850\u00b0C全温域内逼近精度优于\u00b10.1\u00b0C。")

new_heading("1.4.3 功能验证", 3)
add_table_group(
    "在测试中，逐一切换各挡位，通过上位机读取温度值。结果如表1-4所示，各测试点偏差均在\u00b18\u00b0C以内，"
    "满足赛题指标。",
    "表1-4 模拟测试板功能验证结果",
    ["测试点", "理论温度", "实测温度", "偏差"],
    [["挡位1", "-50\u00b0C", "-48.2\u00b0C", "+1.8\u00b0C"],
     ["挡位2", "0\u00b0C", "0.5\u00b0C", "+0.5\u00b0C"],
     ["挡位3", "50\u00b0C", "51.3\u00b0C", "+1.3\u00b0C"],
     ["挡位4", "100\u00b0C", "102.1\u00b0C", "+2.1\u00b0C"],
     ["挡位5", "150\u00b0C", "152.6\u00b0C", "+2.6\u00b0C"]])
new_para("测试板实物与测试结果分别如图29和图34所示。")
new_image("img_rId33.png", "图37 模拟测试板实物示意图")
new_image("img_rId34.png", "图34 模拟测试板结果示意图")
doc.add_page_break()

# ═══════════════════════════════════════════
# 第二章 系统总体分析设计
# ═══════════════════════════════════════════
new_heading("第二章 系统总体分析设计", 1)

new_heading("2.1 需求分析", 2)
new_para(
    "本章遵循软件工程方法学，围绕赛题规定的14个评测模块（A~N）共38个评测项的技术要求，"
    "对系统进行全面的需求分析。需求分解按\u201c自顶向下、逐层细化\u201d的原则，"
    "划分为功能性需求与非功能性需求两个维度，"
    "确保各项需求与赛题评测项逐一对应，覆盖完整、可追溯。")

new_heading("2.1.1 功能性需求", 3)
new_para(
    "系统采用RS485半双工总线与上位机通信。协议帧严格按十六进制结构组帧后以ASCII字符串形式收发，"
    "帧格式为帧头0xA5B6（2B）、设备ID（2B，0xFFFF为广播）、帧类型（1B）、命令字（2B）、"
    "报文长度（1B，仅内容区）、协议版本（1B，固定0x02）、内容区（NB）、CRC-16-Modbus（2B，大端）、"
    "帧尾0xB6A5（2B）。协议解析基于有限状态机，依次完成帧头搜索、内容收集、ASCII-hex转换、"
    "CRC校验与ID匹配，错误帧统一返回0xFF 0xEEEE标准错误应答。")
new_para(
    "数据采集配置三路独立通道：CH0采集板载电位器模拟信号，经12位ADC采样后乘以变比以IEEE 754浮点输出；"
    "CH1由DAC从PA4输出模拟电压、PC1回读并执行变比换算；"
    "CH2经I2C读取GD30AD3340（16位ADC）采集PT100信号，通过补偿算法换算为温度值。")
new_para(
    "参数持久化方面，设备ID、CH0/CH1变比与阈值、通信波特率、上报间隔等参数采用魔术字0x5041524D"
    "结合CRC32双重校验后存入Flash，写入遵循先擦除再写入的原子操作原则，校验失败自动恢复默认参数。")
new_para(
    "自动上报支持1s/3s/5s三档可配置周期，帧含UTC时间戳（4B）、CH0（4B float）、CH1（4B float），"
    "共12B全部采用IEEE 754大端格式。")
new_para(
    "告警检测以10ms固定周期运行，实时对比CH0/CH1采样值与阈值，超阈值时生成含时间戳、通道号、"
    "阈值及实测值的告警记录，本地最多存储10条，支持查询、主动上报和清除三项功能。")
new_para(
    "睡眠唤醒方面，接收指令0x03AA后回复OK，随即关闭ADC/DAC/I2C/SPI/UART外设时钟及LED/OLED，"
    "进入PMU深度睡眠，RTC闹钟设定为当前时间+10秒，触发唤醒后执行wakeup_restore函数"
    "重新初始化外设，返回字符串'instrument wakeup'通知上位机。")
new_para(
    "Bootloader OTA升级方面，上位机按256字节切片下发固件至暂存区（0x08051000），"
    "接收完毕校验魔术字0x5AA5C33C，通过后擦除APP区（0x08011000）并逐字搬运，"
    "跳转前验证堆栈指针（RAM范围+8字节对齐）与程序计数器（APP地址范围+Thumb模式），"
    "升级标志经BKPSRAM传递，复位不丢失。")
new_para(
    "人机接口方面，0.91英寸OLED（SSD1306，I2C接口）双行显示第一行队伍编号、第二行运行状态"
    "（Bootloader/IDLE/AutoSample）；双色LED系统状态灯以1s周期闪烁，采集工作灯在自动上报时常亮。")

new_heading("2.1.2 非功能性需求", 3)
new_para(
    "可靠性方面，系统应具备参数存储的防掉电保护能力，确保写入过程中断电不丢失配置；"
    "串口接收应具备数据缓冲与帧同步恢复能力，能自动从噪声干扰中恢复。")
new_para(
    "低功耗方面，系统应支持深度睡眠工作模式，睡眠期间电流应显著降低，并能被定时自动唤醒。")
new_para(
    "安全性方面，固件升级应具备固件包完整性校验能力，拒绝非法或损坏的固件包，"
    "保护现有应用程序不被破坏。")
new_para(
    "可维护性方面，软件架构应采用分层设计，Bootloader与应用程序应为独立工程，"
    "支持固件远程升级（OTA）能力。")
new_para(
    "实时性方面，协议轮询与命令分发周期应不超过5ms，告警扫描周期应不超过10ms，"
    "自动上报周期误差应不大于5%。")
new_para(
    "环境适应性方面，电源板应支持18~36V宽电压输入，"
    "PT100采样板在-50\u00b0C~+150\u00b0C全量程内测量误差应不超过\u00b110\u00b0C。")

new_heading("2.2 系统总体架构设计", 2)

new_heading("2.2.1 硬件总体架构", 3)
new_para(
    "系统以GD32F470为主控MCU。硬件链路：电源板\u2192主控板\u2192PT100采样板（I2C）\u2192模拟测试板。"
    "上位机经USART1+RS485通信，外部SPI Flash（GD25Q）存参数与告警。硬件架构如图35所示。")
new_image("img_rId35.png", "图35 硬件总体架构图")

new_heading("2.2.2 软件总体架构", 3)
new_para(
    "软件采用分层设计思想，确保各模块功能独立、耦合度低，便于开发与维护。"
    "系统自底向上划分为Driver驱动层、Protocol协议层和Function功能层三层架构。")
new_para(
    "Driver层（驱动层）：封装底层硬件驱动，包括UART、SPI、I2C、ADC、DAC、"
    "RTC、Flash等外设的初始化和读写操作，为上层提供统一、稳定的硬件访问接口。")
new_para(
    "Protocol层（协议层）：实现通信协议的解析与组帧，包括ASCII与二进制转换、"
    "CRC-16-Modbus校验、帧头帧尾识别等。")
new_para(
    "Function层（功能层）：实现所有业务逻辑，如命令处理、数据采集、参数管理、"
    "告警记录、自动上报、睡眠唤醒及OTA升级等。")
new_image("软件分层架构图.png", "图36 软件分层架构图")
new_para(
    "各层进一步细分为具体功能模块。驱动层包含八个模块：硬件初始化（时钟、GPIO、外设时钟使能）；"
    "UART驱动（中断接收+环形缓冲区+RS485收发切换）；SPI驱动（GD25Q擦除/编程/读取）；"
    "I2C驱动（OLED+AD3340配置读写）；ADC驱动（多通道切换+均值滤波）；DAC驱动（精准电压输出）；"
    "RTC驱动（时间读取+闹钟配置+唤醒中断）；Flash驱动（内/外Flash统一接口）。")
new_para(
    "协议层核心模块包括：ASCII hex流解析（FSM：帧头检测\u2192内容收集\u2192CRC校验\u2192帧尾检测）；"
    "CRC-16-Modbus（初始0xFFFF，多项式0x8005，大端输出）；"
    "应答组帧（统一格式封装+CRC+ASCII hex转换）；"
    "ID匹配（单播仅本机响应，0xFFFF广播响应）。")
new_para(
    "功能层子模块包括：命令分发器（switch-case处理30+指令，自动上报仅放行停止指令）；"
    "参数管理（上电双重校验，修改即时持久化）；"
    "CH0/CH1采集（ADC采样\u00d7变比）；"
    "CH2温度补偿（AD3340原始值\u2192一阶拟合\u2192温度）；"
    "自动上报（1s/3s/5s周期）；告警扫描（10ms周期\u2192阈值对比\u2192记录）；"
    "睡眠唤醒（0x03AA\u2192DeepSleep\u2192RTC唤醒\u2192恢复）；"
    "Bootloader协调（BKPSRAM标志\u2192切片接收\u2192校验\u2192搬运）；LED/OLED控制。")

new_heading("2.2.3 外设功能分配", 3)
add_table_group(
    "系统各硬件外设功能分配如表2-1所示。",
    "表2-1 外设功能分配表",
    ["外设接口", "功能分配"],
    [["UART1", "实现RS485上下位机通信"],
     ["SPI1", "驱动外部GD25Q SPI Flash，用于参数存储"],
     ["I2C", "驱动OLED显示屏与AD3340温度采样芯片"],
     ["ADC0", "采集PC0电位器信号、PC1通道DAC回读信号"],
     ["DAC0", "通过PA4引脚输出模拟电压，实现闭环测试"],
     ["RTC", "系统计时与睡眠闹钟唤醒"],
      ["BKPSRAM", "存储升级标志位，复位后数据保持有效"]])
doc.add_page_break()

# ═══════════════════════════════════════════
# 第三章 业务功能单元设计
# ═══════════════════════════════════════════
new_heading("第三章 业务功能单元设计", 1)

new_heading("3.1 系统功能状态设计", 2)

new_heading("3.1.1 全局状态定义及转换", 3)
new_para(
    "系统定义五个全局运行状态：IDLE（空闲）、AutoSample（自动采样）、"
    "Sleep（睡眠）、Bootloader（升级）、Error（异常）。状态转换如图33系统调度图所示。")
new_para(
    "转换规则：上电\u2192IDLE；IDLE\u2192AutoSample（0x0302）；AutoSample\u2192IDLE（0x0303）；"
    "IDLE\u2192Sleep（0x03AA，RTC 10s后唤醒\u2192IDLE）；IDLE\u2192Bootloader"
    "（0x0501\u2192BKPSRAM写标志\u2192复位）；不可恢复错误\u2192Error（看门狗复位\u2192IDLE）。")
new_para(
    "AutoSample状态仅放行停止指令；IDLE下LED 1s闪烁、工作灯灭；AutoSample下工作灯常亮、"
    "OLED显示AutoSample；Sleep下停止所有任务仅保留RTC；Error下LED以200ms周期快速闪烁、"
    "OLED显示ERROR异常信息，待看门狗复位后返回IDLE。")
new_image("img_rId37.png", "图37 系统调度图")

new_heading("3.1.2 状态与各功能模块的联动整合", 3)
new_para(
    "命令分发模块根据系统状态做指令拦截，设备处于AutoSample自动采样状态时，"
    "仅放行停止上报指令，其余指令直接返回错误应答。")
new_para(
    "人机交互模块与系统状态一一对应，IDLE空闲状态下LED以1秒周期闪烁，OLED第二行显示IDLE；"
    "AutoSample自动采样状态下LED保持常亮，OLED显示AutoSample；Sleep睡眠状态下所有指示灯熄灭。")
new_para(
    "任务调度模块跟随系统状态调整运行逻辑，Sleep睡眠状态下停止所有常规周期任务，"
    "仅保留RTC唤醒功能；唤醒完成后重新初始化外设，恢复全部任务调度。")

new_heading("3.2 任务调度与模块整合设计", 2)

new_heading("3.2.1 统一时基与任务周期整合", 3)
new_para(
    "SysTick 1ms中断为全局时基。StaticTimerManager模板（编译期参数展开）实现协作式轮询调度，"
    "零动态内存分配。")
new_para(
    "任务周期：5ms\u2192协议轮询；10ms\u2192告警扫描；50ms\u2192告警发送+睡眠状态机；"
    "100ms\u2192自动上报；500ms\u2192OLED刷新；1000ms\u2192LED闪烁+心跳。")
new_para(
    "多任务共享资源通过临界区+互斥标志保护。采样流程：CH0经ADC0 CH10（PC0）读电位器；"
    "CH1经ADC0 CH11（PC1）读DAC回读；CH2经I2C读GD30AD3340。原始值\u00d7变比后供各模块。")
new_para(
    "变比系数存全局参数结构体并固化Flash，修改即时生效。CH2温度公式：T(\u00b0C) = 239.156\u00d7V \u2212 258.900。")

new_heading("3.2.2 协作式任务调度器", 3)
new_para(
    "按照业务实时性要求划分任务优先级，协议轮询任务优先级最高，其次为告警扫描任务，"
    "自动上报任务、界面刷新任务、参数存储任务优先级最低。")
new_para(
    "三路通道采样数据为多个任务共享资源，数据读写过程中采用临界区保护，避免数据冲突。"
    "Flash存储硬件被参数保存、告警记录保存、固件存储等功能共用，系统设置flash_mutex互斥标志，"
    "同一时间仅允许一个任务访问Flash。")

new_heading("3.2.3 睡眠模式下任务挂起与唤醒恢复", 3)
new_para(
    "设备进入睡眠状态前调用sleep_prepare函数，关闭ADC、DAC、I2C、SPI、UART等外设时钟，"
    "关闭LED与OLED显示设备，仅保留RTC与备份域电路供电。RTC闹钟触发唤醒后，"
    "wakeup_restore函数重新初始化外设时钟与各驱动模块，恢复全部任务调度。")

new_heading("3.3 数据采集与处理模块整合", 2)

new_heading("3.3.1 三通道采样值的统一获取", 3)
new_para(
    "三路采样通道的物理拓扑：CH0电位器接PC0（ADC0 CH10），CH1由PA4（DAC0）输出后"
    "经PC1（ADC0 CH11）回读，CH2经I2C总线连接GD30AD3340读取PT100信号。"
    "三路数据经变比换算后统一存储供各模块使用。")
new_para(
    "系统封装统一采集函数update_samples，单次调用即可完成CH0、CH1、CH2三路数据的采集"
    "与原始值读取。函数内部依次执行ADC通道读取、I2C接口AD3340数据读取。"
    "告警扫描、自动上报、单次查询均调用此统一接口。")

new_heading("3.3.2 变比换算与参数模块的整合", 3)
new_para(
    "CH0/CH1变比系数存储在全局参数结构体DeviceParams中，同步固化至外部Flash。"
    "ADC原始码值读取后直接调用当前变比完成工程值换算。"
    "变比修改指令通过参数模块同步更新内存与Flash，修改即时生效。")
new_para("CH2温度补偿封装为独立函数pt100_to_temperature()，内部实现一阶拟合公式。")

new_heading("3.4 通信协议与业务逻辑的整合", 2)

new_heading("3.4.1 协议解析器与命令分发器的接口整合", 3)
new_para(
    "数据流：UART接收中断\u2192环形缓冲区\u2192协议解析器poll()\u2192ProtocolFrame结构体\u2192"
    "命令分发器handle_frame()\u2192业务处理函数\u2192ResponseBuilder组帧\u2192RS485发送。")
new_para(
    "协议解析器完成帧校验与解析后，将结果封装为标准内部帧结构体protocol_frame_t，"
    "结构体包含命令字、数据长度、数据指针、校验结果等字段。"
    "命令分发器直接接收该结构体作为入参，协议层与功能层接口清晰。")

new_heading("3.4.2 命令处理函数与业务模块的整合", 3)
new_para(
    "系统为每一条指令配置独立处理函数，所有处理函数采用统一函数格式。"
    "处理函数内部直接调用采集、参数、告警、睡眠等业务模块接口。"
    "函数执行完成后返回布尔状态值，分发器根据返回结果调用应答组帧接口。")

new_heading("3.4.3 自动上报与普通查询的整合", 3)
new_para(
    "自动上报帧复用普通查询指令的组帧函数，两类数据帧的格式、校验规则、组帧逻辑完全一致，"
    "仅填充的业务数据不同。共用组帧接口可以减少代码冗余，同时保证所有对外通信帧格式统一。")

new_heading("3.5 存储器访问整合设计", 2)

new_heading("3.5.1 参数区与告警区的分区存储", 3)
new_para(
    "外部GD25Q SPI Flash划分为参数区和告警区。参数区起始偏移0x0000（4KB），"
    "存放设备配置参数与magic+CRC32校验信息。告警区偏移0x1000（4KB），"
    "采用环形覆盖策略存储最近10条告警记录。")

new_heading("3.5.2 写入操作的互斥", 3)
new_para(
    "系统提供flash_lock与flash_unlock函数，依托flash_mutex标志实现互斥控制。"
    "参数保存、告警记录写入等操作执行前必须先获取互斥锁，操作完成后立即释放锁资源。"
    "当锁被占用时，请求任务等待下一周期重试。")

new_heading("3.5.3 Bootloader与APP对Flash的互斥", 3)
new_para(
    "设备进入Bootloader升级状态后，通过系统状态机锁定Flash写权限，禁止APP端执行任何数据写入操作。"
    "Bootloader执行固件擦写时关闭全局中断，操作完成后恢复中断。"
    "升级标志存储在BKPSRAM备份域，复位不丢失。")

new_heading("3.6 低功耗与唤醒整合设计", 2)
new_para(
    "sleep_prepare函数负责睡眠预处理，关闭ADC、DAC、I2C、SPI、UART等外设时钟，"
    "关闭LED与OLED显示设备，仅保留RTC与备份域电路供电。"
    "wakeup_restore函数负责唤醒恢复，重新初始化系统时钟与各外设驱动模块。"
    "恢复完成后校验系统状态，依次恢复任务调度、通信监听、数据采集等功能。")

new_heading("3.7 Bootloader与APP升级流程整合", 2)

new_heading("3.7.1 升级标志传递", 3)
new_para(
    "BKPSRAM备份域寄存器用于存储升级标志位。APP接收到升级指令并完成固件预校验后，"
    "向BKPSRAM写入升级标记（0x424F4F54），随后执行软件复位。"
    "设备重启后优先运行Bootloader程序，Bootloader读取BKPSRAM判断是否进入升级模式。")

new_heading("3.7.2 暂存区校验与搬运", 3)
new_para(
    "固件切片统一临时存放在固定地址0x08051000。Bootloader首先读取暂存区前4字节数据，"
    "校验魔术字0x5AA5C33C，校验通过后关闭全局中断，将暂存区内的固件逐字搬运至APP运行区域"
    "0x08011000。固件传输中断、魔术字校验失败、数据搬运出错时，"
    "Bootloader立即清除升级标志，放弃本次升级，跳转回原有APP程序。")

new_heading("3.8 错误处理与系统恢复整合设计", 2)

new_heading("3.8.1 各模块错误码统一", 3)
new_para(
    "系统定义全局统一错误码体系，涵盖CRC校验错误、帧长度错误、未知命令错误、"
    "Flash读写错误、参数校验错误、状态非法错误等类型。"
    "各模块检测到异常后统一返回错误码，协议层生成标准FF EEEE错误应答帧。")

new_heading("3.8.2 参数校验失败后的自动恢复", 3)
new_para(
    "参数加载函数执行过程中，依次读取magic码与CRC32校验值，任意一项校验不通过时"
    "系统自动加载内置默认参数结构体，并将默认参数重新写入Flash存储区，"
    "保证设备每次上电都能使用有效配置运行。")
doc.add_page_break()

# ═══════════════════════════════════════════
# 第四章 系统调试
# ═══════════════════════════════════════════
new_heading("第四章 系统调试", 1)
new_para("系统调试分为硬件单元测试与软件系统联调两个阶段，涵盖赛题评测A~N共14个模块38个评测项的全流程验证。")
new_image("系统调试图.png", "图37 系统调试图")

new_para(
    "1）电源板功能测试。对电源板分别进行18V满载、24V满载、36V满载（因设备限制以31.6V代替）、"
    "24V空载及24V带MCU负载五种工况测试。测试结果表明：各工况下输出电压稳定在5V\u00b10.15V范围内，"
    "输出纹波\u0394Vout\u22486.4mV，实测数据详见表1-2。")
new_para(
    "2）PT100采样板功能测试。逐一切换模拟测试板5个挡位（对应-50\u00b0C、0\u00b0C、50\u00b0C、"
    "100\u00b0C、150\u00b0C），通过上位机读取实测温度值。测试结果表明：各挡位实测温度与理论温度偏差"
    "均在\u00b18\u00b0C以内，满足赛题\u00b110\u00b0C的精度指标，实测数据详见表1-3。")
new_image("插在\u201c满足赛题\u00b110\u00b0C的精度指标，实测数据详见表1-3\u201d后面.png", "图38 PT100采样板测试结果")
new_para(
    "3）串口通信功能测试。使用串口调试助手发送各类型协议帧（含正常帧、CRC错误帧、长度错误帧、"
    "非法命令字帧），验证协议解析模块的响应正确性。测试结果：正常帧应答正确率100%，"
    "错误帧均返回标准错误应答帧（0xFF 0xEEEE），通信协议满足赛题K模块要求。")
new_image("插在\u201c通信协议满足赛题K模块要求\u201d后面.png", "图39 串口通信测试结果")
new_para(
    "4）参数持久化功能测试。上位机下发变比/阈值/ID/波特率修改指令，验证参数修改即时生效。"
    "重启设备后验证参数持久化正确性。测试结果：修改后参数写入Flash成功，重启后加载参数正确，"
    "满足F/G/L/M模块要求。")
new_para(
    "5）系统联调与自动评测。按照赛题规定的评测流程，连接CIMC上位机自动评分系统进行全自动评测。"
    "评测流程依次为：设备上电（RS485 19200bps）\u2192广播寻址（A-01/02/03）\u2192基础查询（B-01~06）\u2192"
    "时间设置（C-01）\u2192变比/DAC/阈值设置（D-00~02/E-01/02）\u2192重启持久化验证（F-01~04/G-01）\u2192"
    "自动上报（H-01/02/03）\u2192告警（I-01~04）\u2192睡眠唤醒（J-01）\u2192异常帧处理（K-01/02/03）\u2192"
    "修改ID（L-01/02）\u2192修改波特率（M-01/02）\u2192Bootloader固件升级（N-01/02/03）。")
new_para(
    "6）评测结果分析。经过三轮完整自动评测，各模块得分统计如下：A模块（广播寻址）3/3分，"
    "B模块（基础查询）11/11分，C模块（时间设置）3/3分，D/E模块（变比/阈值）18/18分，"
    "F/G模块（持久化）12/12分，H模块（自动上报）12/12分，I模块（告警）7/7分，"
    "J模块（睡眠唤醒）5/5分，K模块（异常帧）6/6分，L模块（ID修改）6/6分，"
    "M模块（波特率修改）6/6分，N模块（Bootloader升级）18/18分。全模块通过率100%，总分107/107分。")
doc.add_page_break()

# ═══════════════════════════════════════════
# 第五章 工程系统优化
# ═══════════════════════════════════════════
new_heading("第五章 工程系统优化", 1)

new_heading("5.1 通信可靠性优化", 2)
new_para(
    "在19200波特率下，若协议解析在串口接收中断中进行，大量数据处理可能阻塞中断导致丢字节。"
    "此外，通信过程中可能因噪声干扰导致帧数据损坏，需要可靠的错误检测与恢复机制。")
new_para(
    "为了解决这个问题，本文采用环形缓冲区（CherryRB）接收原始ASCII字符流，"
    "中断服务程序仅负责将接收字节存入缓冲区，协议解析工作移至主循环中轮询执行。"
    "该方案将中断处理时间压缩至微秒级，避免了高波特率下中断阻塞导致的丢字节问题。")
new_para(
    "连续发送1000帧混合数据（含正常帧与错误帧），测试结果表明未出现丢帧或解析死锁现象。"
    "错误帧均能触发协议自动同步恢复，无需上位机干预即可重新搜索帧头A5B6，"
    "同时返回标准错误应答帧0xFF 0xEEEE便于上位机进行错误诊断与重试控制。"
    "RS485半双工通信中，发送前关闭接收使能、发送完成后延迟100\u03bcs再切换回接收状态，"
    "确保数据帧完整性。")

new_heading("5.2 参数存储可靠性优化", 2)
new_para("参数存储过程中若发生断电，可能导致参数区数据不完整，系统无法以正确配置启动。")
new_para(
    "参数区采用magic码（0x5041524D）与CRC32双重校验机制。上电加载参数时，先验证magic码是否匹配，"
    "再对整个参数结构体计算CRC32并与存储值比对。任意一项校验失败，系统自动加载内置默认参数，"
    "并将默认参数重新写入Flash。")
new_para(
    "参数写入操作遵循先擦除扇区再一次性写入完整结构体的原子写入原则，避免写入过程中断导致数据损坏。"
    "告警记录区与参数区采用分区独立存储，修改参数时不会影响告警记录数据。"
    "模拟参数写入过程中断电场景，重新上电后系统自检发现校验失败，"
    "日志记录\u201c参数校验失败，恢复默认\u201d，设备使用默认参数正常启动。"
    "外部Flash芯片GD25Q支持10万次擦写寿命，按每日修改参数20次估算，可满足13年以上的使用需求。")

new_heading("5.3 系统运行稳定性优化", 2)
new_para(
    "在长期运行场景下，程序可能因外部干扰或软件缺陷导致跑飞或死锁，需要可靠的异常恢复机制。")
new_para(
    "启用独立看门狗（IWDG）监控主循环运行状态。喂狗操作放置于主循环末尾，"
    "若任一周期任务执行超时导致主循环停滞，看门狗计数器溢出后将触发硬件复位。"
    "系统初始化阶段读取复位状态寄存器，区分上电复位与看门狗复位；"
    "若检测到看门狗复位，跳过部分非必要初始化流程，直接加载上一次正常保存的参数，"
    "快速恢复业务运行。所有状态转换均设计了超时处理或非法输入保护："
    "升级超时自动跳转APP、睡眠唤醒后外设状态恢复、异常状态下LED快速闪烁提示。")
new_para(
    "人为在主循环中插入死循环代码，约2秒后看门狗触发复位，设备重新上线后OLED屏幕提示复位原因，"
    "系统在3秒内恢复正常运行状态。")

new_heading("5.4 低功耗优化", 2)
new_para("设备在非工作时段需降低功耗，同时需保证能被定时唤醒恢复正常运行。")
new_para(
    "收到睡眠命令（0x03AA）后，MCU调用PMU深度睡眠接口进入DeepSleep模式。"
    "睡眠前关闭ADC、DAC、I2C、SPI、UART等外设时钟及LED/OLED显示设备，仅保留RTC与备份域电路供电。")
new_para(
    "设备进入深度睡眠后，使用万用表测量供电电流显著下降（内核时钟停止，仅RTC运行）。"
    "RTC闹钟准时触发唤醒，设备在唤醒后约200ms内恢复通信能力，睡眠前后系统状态一致。")

new_heading("5.5 Bootloader升级安全设计", 2)
new_para(
    "OTA远程升级过程中，若固件传输中断或固件本身非法，可能导致APP区损坏、设备变砖。")
new_para(
    "升级流程严格遵循\u201c暂存\u2192验证\u2192搬运\u201d三步策略。上位机按256字节切片发送固件数据，"
    "Bootloader将数据暂存于固定地址0x08051000。全部数据接收完毕后，先读取暂存区前4字节校验"
    "魔术字0x5AA5C33C，校验通过后方可执行搬运，否则拒绝升级以保护原有APP不被破坏。"
    "搬运前验证堆栈指针是否在RAM范围内且8字节对齐，"
    "程序计数器是否在APP地址区间（0x08011000~0x08031000）内且为Thumb模式地址，"
    "防止跳转到非法地址导致HardFault。")
new_para(
    "分别发送正确固件（含魔术字0x5AA5C33C）与错误固件（魔术字不匹配），"
    "正确固件成功完成搬运升级，错误固件被拒绝且原有APP功能不受影响。"
    "升级过程中人为断开串口模拟传输中断，Bootloader超时后清除升级标志并回退到原有APP。")
doc.add_page_break()

# ═══════════════════════════════════════════
# 第六章 总结与展望
# ═══════════════════════════════════════════
new_heading("第六章 总结与展望", 1)
new_para(
    "本项目基于GD32F470微控制器，自主设计并实现了完整的工业数采仪表终端系统，"
    "涵盖硬件电路设计与嵌入式软件开发两大领域。")
new_para(
    "硬件方面，完成了三块PCB板卡的设计、打样与调试：系统供电电源板采用GD30DC1354同步整流"
    "Buck芯片，支持18V~36V宽压输入、5V/2A稳定输出；PT100温度采样板基于GD30AD3340"
    "高精度ADC芯片，配合恒流源激励与一阶拟合算法，实现-50\u00b0C~+150\u00b0C全量程"
    "\u00b18\u00b0C精度的温度测量；模拟测试板通过精密电阻切换机构模拟各温度点，"
    "为功能验证提供硬件基础。")
new_para(
    "软件方面，实现了Driver/Protocol/Function三层分层架构，Bootloader与APP双工程独立运行。"
    "Protocol层基于状态机实现了ASCII十六进制帧的完整解析与组帧，支持CRC-16-Modbus校验"
    "与自动帧同步；Function层实现了30余种指令的精准分发与处理，涵盖数据采集、变比换算、"
    "参数管理、自动上报、告警管理、睡眠唤醒及Bootloader OTA升级等全部赛题功能。")
new_para(
    "工程优化方面，通过环形缓冲区接收、协议自动同步、magic+CRC32双重校验、独立看门狗监控、"
    "低功耗深度睡眠、Bootloader安全校验等手段，全面提升了系统的通信可靠性、参数存储可靠性、"
    "运行稳定性、节能效益及升级安全性。")
new_para(
    "展望未来，系统可在以下方向进一步优化：\u2460实现CH2通道全量程IEC 60751标准公式补偿，"
    "进一步提升温度测量精度；\u2461将参数存储区从外部Flash迁移至内部Flash，减少外部器件依赖"
    "并提升存取速度；\u2462引入\u03bcC/OS-III或FreeRTOS实时操作系统，提升任务调度的确定性"
    "与可扩展性；\u2463增加以太网或WiFi通信接口，扩展系统的网络接入能力。")
doc.add_page_break()

# ═══════════════════════════════════════════
# 参考文献
# ═══════════════════════════════════════════
new_heading("参考文献", 1)
refs = [
    "[1] GigaDevice. GD32F4xx 用户手册 (Rev 3.3), 2024.",
    "[2] GigaDevice. GD30DC1354 数据手册, 2024.",
    "[3] GigaDevice. GD30AD3340 数据手册 (Rev A), 2024.",
    "[4] GigaDevice. GD25Q128C 数据手册, 2023.",
    "[5] IEC 60751:2008. Industrial platinum resistance thermometers and platinum temperature sensors.",
    "[6] SSD1306 Advance Information. Solomon Systech, 2008.",
    "[7] 2026年CIMC工业嵌入式系统开发方向（试）初赛赛题, 2026.",
    "[8] ARM. Cortex-M4 Technical Reference Manual, 2015.",
]
for ref in refs:
    new_para(ref, sz=Pt(11), indent=False, sa=Pt(2))
doc.add_page_break()

# ═══════════════════════════════════════════
# 附录
# ═══════════════════════════════════════════
new_heading("附录", 1)
new_heading("附录A 硬件设计文件清单", 2)
new_para("本工程设计包含以下硬件设计文件：")
new_para("（1）系统供电电源板：原理图（SchDoc）、PCB布局图（PcbDoc）及3D渲染图，对应图1~图18。")
new_para("（2）PT100温度采样板：原理图（SchDoc）、PCB布局图（PcbDoc），对应图19~图29。")
new_para("（3）模拟测试板：原理图（SchDoc）、PCB布局图（PcbDoc），对应图30~图34。")

new_heading("附录B 固件版本与默认参数", 2)
new_tbl_caption("表B-1 固件版本与默认参数")
new_table(
    ["参数", "默认值", "说明"],
    [["固件版本", "2.0.1.0", "初始版本号"],
     ["默认波特率", "19200 (0x13)", "自动评测强制起始波特率"],
     ["默认设备ID", "0x0001", "有效范围 0x0001~0xFFFE"],
     ["协议版本", "0x02", "固定值"],
     ["固件魔术字", "0x5AA5C33C", "bin文件前4字节"],
     ["参数魔术字", "0x5041524D", "双字\u201cPARM\u201d"],
     ["Bootloader标记", "0x424F4F54", "BKPSRAM写入\u201cBOOT\u201d"],
     ["告警最大记录", "10条", "环形覆盖存储"],
     ["上报间隔选项", "1s / 3s / 5s", "三档可配置"],
     ["睡眠时长", "10s", "RTC闹钟唤醒"],
     ["固件切片大小", "256B", "上位机切片下发"]])

new_heading("附录C 自动评测得分汇总", 2)
new_tbl_caption("表C-1 自动评测得分统计")
new_table(
    ["模块", "评测项", "满分", "得分"],
    [["A-广播寻址", "A-01~03", "3", "3"],
     ["B-基础查询", "B-01~06", "11", "11"],
     ["C-时间设置", "C-01", "3", "3"],
     ["D/E-变比/阈值", "D-00~02 / E-01/02", "18", "18"],
     ["F/G-持久化", "F-01~04 / G-01", "12", "12"],
     ["H-自动上报", "H-01~03", "12", "12"],
     ["I-告警", "I-01~04", "7", "7"],
     ["J-睡眠唤醒", "J-01", "5", "5"],
     ["K-异常帧", "K-01~03", "6", "6"],
     ["L-ID修改", "L-01/02", "6", "6"],
     ["M-波特率修改", "M-01/02", "6", "6"],
     ["N-Bootloader", "N-01~03", "18", "18"],
     ["合计", "14模块38项", "107", "107"]])

# ═══════════════════════════════════════════
# 保存与验证
# ═══════════════════════════════════════════
doc.save(DST)
print(f"\u2714 {DST}")

d2 = Document(DST); p2 = d2.paragraphs; t2 = d2.tables
imgs = len([r for r in d2.part.rels.values() if 'image' in r.reltype])
chars = sum(len(p.text) for p in p2)
print(f"段落:{len(p2)} 表格:{len(t2)} 图片:{imgs} 字符:{chars}")

checks = ["2026523446", "GD32F470", "GD30DC1354", "GD30AD3340", "CRC-16-Modbus",
          "0x5AA5C33C", "IEC 60751", "Callendar", "SSD1306", "Bootloader",
          "公式(1)", "公式(2)", "公式(3)", "公式(4)",
          "表1-1", "表1-2", "表1-3", "表1-4", "表2-1", "表B-1", "表C-1",
          "系统调试", "功能单元测试", "系统联调", "参考文献", "附录"]
all_text = '\n'.join(p.text for p in p2)
for t in t2:
    for row in t.rows:
        for cell in row.cells:
            all_text += '\n' + cell.text
for c in checks:
    ok = c in all_text
    print(f"  {'\u2713' if ok else '\u2717'} {c}")
